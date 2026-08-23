"""验证知识文档解析器的格式识别、结构保留和恶意输入边界。"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from pypdf import PdfWriter

from aicare_agent_service.rag.contracts import KnowledgeMetadata, RawKnowledgeDocument
from aicare_agent_service.rag.errors import RagError, RagErrorCode
from aicare_agent_service.rag.parsers import parse_document

FIXTURES = Path(__file__).parents[1] / "fixtures" / "rag"


def _metadata() -> KnowledgeMetadata:
    return KnowledgeMetadata(
        tenant_id="tenant-1",
        knowledge_base_id="kb-1",
        document_id="doc-1",
        version=1,
        language="zh-CN",
        category="POLICY",
    )


def _raw(media_type: str, content: bytes, file_name: str) -> RawKnowledgeDocument:
    return RawKnowledgeDocument(
        metadata=_metadata(),
        file_name=file_name,
        media_type=media_type,
        source_uri="https://kb.example/doc-1",
        content=content,
    )


def _docx_bytes(*, external_link: bool = False, macro: bool = False) -> bytes:
    stream = BytesIO()
    document = Document()
    document.add_heading("账号交付", level=1)
    document.add_paragraph("成品账号由人工审核后交付。")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "类型"
    table.rows[0].cells[1].text = "时效"
    document.save(stream)
    if not external_link and not macro:
        return stream.getvalue()

    output = BytesIO()
    with (
        ZipFile(BytesIO(stream.getvalue())) as source,
        ZipFile(output, "w", compression=ZIP_DEFLATED) as target,
    ):
        for entry in source.infolist():
            payload = source.read(entry.filename)
            if external_link and entry.filename == "word/_rels/document.xml.rels":
                payload = payload.replace(
                    b"</Relationships>",
                    b'<Relationship Id="evil" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" Target="https://attacker.invalid" TargetMode="External"/></Relationships>',
                )
            target.writestr(entry, payload)
        if macro:
            target.writestr("word/vbaProject.bin", b"macro")
    return output.getvalue()


def _pdf_bytes(*, encrypted: bool = False, pages: int = 1) -> bytes:
    stream = BytesIO()
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=300, height=300)
    if encrypted:
        writer.encrypt("secret")
    writer.write(stream)
    return stream.getvalue()


def _text_pdf_bytes() -> bytes:
    """构造带Helvetica文字层的固定最小PDF样例。"""
    content = b"BT /F1 12 Tf 50 250 Td (Steam gift delivery policy) Tj ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 300] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream",
    ]
    payload = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, value in enumerate(objects, start=1):
        offsets.append(len(payload))
        payload.extend(f"{index} 0 obj\n".encode() + value + b"\nendobj\n")
    xref = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    payload.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        payload.extend(f"{offset:010d} 00000 n \n".encode())
    payload.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode()
    )
    return bytes(payload)


@pytest.mark.parametrize(
    ("fixture_name", "media_type", "expected"),
    [
        ("policy.txt", "text/plain", "Steam 礼物"),
        ("policy.md", "text/markdown", "支付成功后自动交付"),
        ("policy.html", "text/html", "未展示的 CDK"),
    ],
)
def test_parse_text_formats_from_bytes(fixture_name: str, media_type: str, expected: str) -> None:
    sections = parse_document(
        _raw(media_type, (FIXTURES / fixture_name).read_bytes(), fixture_name)
    )

    assert expected in "\n".join(section.text for section in sections)


def test_markdown_preserves_heading_list_and_code_boundaries() -> None:
    sections = parse_document(
        _raw("text/markdown", (FIXTURES / "policy.md").read_bytes(), "policy.md")
    )

    assert any(section.title_path == ("交付政策", "CDK") for section in sections)
    rendered = [section.text for section in sections]
    assert "- 支付成功后自动交付。\n- 请在 Steam 客户端激活。" in rendered
    assert "```text\n激活路径:游戏 > 在 Steam 上激活产品\n```" in rendered


def test_html_removes_active_content_without_fetching_external_resources() -> None:
    sections = parse_document(
        _raw("text/html", (FIXTURES / "policy.html").read_bytes(), "policy.html")
    )

    rendered = "\n".join(section.text for section in sections)
    assert "window.evil" not in rendered
    assert "attacker.invalid" not in rendered
    assert "类型 | 处理方式" in rendered
    assert any(section.title_path == ("退款政策", "申请条件") for section in sections)


def test_docx_preserves_heading_and_table_rows() -> None:
    sections = parse_document(
        _raw(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _docx_bytes(),
            "policy.docx",
        )
    )

    assert any(section.title_path == ("账号交付",) for section in sections)
    assert "类型 | 时效" in [section.text for section in sections]


def test_text_pdf_preserves_page_location_and_extracts_text_layer() -> None:
    sections = parse_document(_raw("application/pdf", _text_pdf_bytes(), "policy.pdf"))

    assert sections[0].title_path == ("第 1 页",)
    assert sections[0].text == "Steam gift delivery policy"


@pytest.mark.parametrize(
    "document",
    [
        _raw("text/plain", b"%PDF-1.7\nnot-text", "fake.txt"),
        _raw("application/pdf", b"not-a-pdf", "fake.pdf"),
        _raw("text/plain", b"safe\x00hidden", "control.txt"),
        _raw(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _docx_bytes(macro=True),
            "macro.docm",
        ),
        _raw(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _docx_bytes(external_link=True),
            "external.docx",
        ),
    ],
)
def test_parser_rejects_spoofed_or_active_documents(document: RawKnowledgeDocument) -> None:
    with pytest.raises(RagError) as exc_info:
        parse_document(document)

    assert exc_info.value.code == RagErrorCode.UNSUPPORTED_FORMAT


def test_parser_rejects_docx_zip_bomb() -> None:
    stream = BytesIO()
    with ZipFile(stream, "w", compression=ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", b"<Types/>")
        archive.writestr("word/document.xml", b"A" * (8 * 1024 * 1024))

    with pytest.raises(RagError) as exc_info:
        parse_document(
            _raw(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                stream.getvalue(),
                "bomb.docx",
            )
        )

    assert exc_info.value.code == RagErrorCode.DOCUMENT_TOO_LARGE


@pytest.mark.parametrize("content", [_pdf_bytes(), _pdf_bytes(encrypted=True)])
def test_scanned_or_encrypted_pdf_is_not_silently_indexed(content: bytes) -> None:
    with pytest.raises(RagError) as exc_info:
        parse_document(_raw("application/pdf", content, "scan.pdf"))

    assert exc_info.value.code == RagErrorCode.NO_EXTRACTABLE_TEXT


def test_pdf_page_limit_is_enforced_before_text_extraction() -> None:
    with pytest.raises(RagError) as exc_info:
        parse_document(_raw("application/pdf", _pdf_bytes(pages=301), "large.pdf"))

    assert exc_info.value.code == RagErrorCode.DOCUMENT_TOO_LARGE
